import io
import json

import pdfplumber
from celery import shared_task
from docx import Document
from openai import OpenAI

client = OpenAI()


def extract_docx_text(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    text = []

    for para in doc.paragraphs:
        if para.text.strip():
            text.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    text.append(cell_text)

    for section in doc.sections:
        header = section.header
        footer = section.footer

        for para in header.paragraphs:
            if para.text.strip():
                text.append(para.text)

        for para in footer.paragraphs:
            if para.text.strip():
                text.append(para.text)

    return "\n".join(text)


@shared_task
def generate(file_bytes: bytes, filename: str, program_outcomes_str: str):

    text_content = ""
    suffix = filename.split(".")[-1].lower()

    if suffix == "docx":
        text_content = extract_docx_text(file_bytes)
    elif suffix == "pdf":
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content += page_text + "\n"
    else:
        return {"error": "Unsupported file type"}

    prompt = f"""
You are a university course coordinator. I will provide a syllabus for a course and a list of Program Outcomes (POs).

Syllabus:
{text_content}

Program Outcomes:
{program_outcomes_str}

Instructions: 
1. Read the syllabus carefully.
2. Generate Learning Outcomes (LOs) strictly based on the course aims, course outputs, and course content. Do NOT invent outcomes that are not mentioned in the syllabus.
3. For each LO, assign 1–4 Program Outcomes (POs) with a weight 1–5.
4. Identify only the measurable evaluation components explicitly listed in the syllabus: Midterm, Project, Assignment, Final.
5. Assign relevance score (1–5) for each LO to each assessment it maps to.
6. Output **only raw JSON** with this structure:
{{
  "LOs": [
    {{ "LO": "LO1", "description": "...", "POs": {{ "PO1": 5, "PO3": 4 }} }},
    ...
  ],
  "Assessments": {{
    "Midterm": {{ "LO1": 5, "LO2": 4 }},
    "Assignment": {{ "LO1": 4, "LO2": 5 }},
    "Project": {{ "LO2": 4, "LO3": 5 }},
    "Final": {{ "LO1": 5, "LO3": 4 }}
  }}
}}
7. Do NOT include markdown, code blocks, or any explanations.
"""

    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {
                "role": "system",
                "content": "You are an expert university course coordinator.",
            },
            {"role": "user", "content": prompt},
        ],
    )

    result_text = response.choices[0].message.content

    try:
        result_json = json.loads(result_text)
    except json.JSONDecodeError:
        import re

        cleaned = re.sub(r"```(?:json)?|```", "", result_text).strip()
        try:
            result_json = json.loads(cleaned)
        except json.JSONDecodeError:
            result_json = {"raw_text": result_text}

    return result_json


@shared_task
def generate_program_suggestions(program_payload: dict):
    """
    program_payload örneği:
    {
        "program_name": "Computer Science Engineering",
        "program_outcomes": { ... },
        "courses": [ ... ]
    }
    """
    prompt = f"""
You are an expert university program coordinator. I will provide you with a program's name, a list of Program Outcomes (POs) with descriptions, and the courses in the program with their assessments and average scores.

Program Name:
{program_payload['program_name']}

Program Outcomes:
{json.dumps(program_payload['program_outcomes'], indent=2)}

Courses:
{json.dumps(program_payload['courses'], indent=2)}

Instructions:
1. Analyze the program outcomes and the courses' assessment scores.
2. Identify areas where the program can be improved based on low PO scores or course performance.
3. Provide actionable suggestions for curriculum improvement, course redesign, or assessment strategies to improve PO achievements.
4. Output only JSON with the following structure:

{{
  "Suggestions": [
    {{"PO": "PO1", "course": "CSE311", "issue": "Low assessment scores", "recommendation": "Increase practical assignments and provide more lab exercises"}}
  ]
}}

Do not include any markdown, explanations, or text outside the JSON.
"""

    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {
                "role": "system",
                "content": "You are an expert university program coordinator.",
            },
            {"role": "user", "content": prompt},
        ],
    )

    result_text = response.choices[0].message.content

    try:
        result_json = json.loads(result_text)
    except json.JSONDecodeError:
        import re

        cleaned = re.sub(r"```(?:json)?|```", "", result_text).strip()
        try:
            result_json = json.loads(cleaned)
        except json.JSONDecodeError:
            result_json = {"raw_text": result_text}

    return result_json
